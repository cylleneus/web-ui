import re

import docx
from html3.html3 import HTML

import cylleneus.settings
from cylleneus.utils import slugify


def as_docx(search, filename):
    if filename == "":
        filename = slugify(search.query, allow_unicode=False)
    else:
        filename = slugify(filename, allow_unicode=False)

    if search.results:
        doc = docx.Document()
        doc.add_heading(search.query, 0)
        doc.add_heading(
            f"{search.count_matches} matches in {search.count_documents} documents "
            f"({search.start_dt.strftime(cylleneus.settings.LONG_DATE_FORMAT)})",
            2,
        )

        for result in search.results:
            html = result.html
            author = re.search(r'id="author">(.+?)</div>', html)
            title = re.search(r'id="title">(.+?)</div>', html)
            reference = re.search(r'id="reference">(.+?)</div>', html)
            text = re.search(r'id="text">(.+?)</div>', html, flags=re.DOTALL)

            h = doc.add_heading(level=1)
            h.add_run(f"{author.group(1)}, ")
            r = h.add_run(f"{title.group(1)} ")
            r.font.italic = True
            h.add_run(f"{reference.group(1)}")

            p = doc.add_paragraph()
            for run in re.finditer(
                    r"<(\w+?) .*?>(.*?)</\1>", text.group(1), flags=re.DOTALL
            ):
                if run.group(1) == "match":
                    for t in run.group(2).split():
                        if re.match(r"<em>.*?</em>", t):
                            em = re.match(r"<em>(.*?)</em>", t).group(1)
                            r = p.add_run(em + " ")
                            r.font.bold = True
                        elif t.startswith("<em>"):
                            r = p.add_run(re.sub(r"<em>", "", t) + " ")
                            r.font.bold = True
                        elif t.startswith("</em>"):
                            r = p.add_run(re.sub(r"</em>", "", t) + " ")
                            r.font.bold = False
                        elif t == "\n":
                            r.add_break()
                        else:
                            r = p.add_run(t + " ")
                            r.font.bold = None
                else:
                    p.add_run(run.group(2))
        doc.save(f"{filename}.docx")


def as_html(highlights):
    """Return formatted HTML elements for the given results object."""

    for href in highlights:
        htm = HTML()

        h = htm.div()
        h.div(href.author, klass="h5 card-title font-weight-bold exportable", id="author")
        h.div(href.title, klass="h5 card-title font-italic exportable", id="title")
        if href.reference:
            h.div(
                href.reference, klass="h6 card-subtitle mb-2 text-muted pt-1 exportable", id="reference"
            )

        d = htm.div(klass="pt-1 exportable", id="text")
        # Process post-match context
        pre_texts = re.finditer(r"<pre>(.*?)</pre>(\n\n)?", href.text, flags=re.DOTALL)
        if pre_texts:
            for pre_text in pre_texts:
                if pre_text.group(2):
                    klass = "card-text d-block"
                else:
                    klass = "card-text"
                d.span(pre_text.group(1), klass=f"{klass}")

        # Process matched text
        match_texts = re.finditer(r"<match>(.*?)</match>(\n\n)?", href.text, flags=re.DOTALL)
        if match_texts:
            for match_text in match_texts:
                if match_text.group(2):
                    klass = "card-text d-block"
                else:
                    klass = "card-text"

                text = re.sub(r" </em>", "</em> ", match_text.group(1))
                tokens = text.split()

                buffer = []
                for token in tokens:
                    if re.match(r"<em>(.*?)</em>", token):
                        if buffer:
                            d.span(" ".join(buffer), klass=f"{klass}")
                            buffer.clear()
                        em = re.sub(r"<em>(.*?)</em>", r" \1 ", token, flags=re.DOTALL)
                        d.span(em, klass=f"{klass} text-primary font-weight-bold")
                    else:
                        buffer.append(token)

        # Process post-match context
        post_texts = re.finditer(r"<post>(.*?)</post>(\n\n)?", href.text, flags=re.DOTALL)
        if post_texts:
            for post_text in post_texts:
                if post_text.group(2):
                    klass = "card-text d-block"
                else:
                    klass = "card-text"
                d.span(post_text.group(1), klass=f"{klass}")

        # TODO: Add links to full text
        # if href.urn != "":
        #     d = htm.div(klass="float-right")
        #     p = d.p('more...')
        #     p.a(href=href.urn, klass="card-link exportable")

        yield str(htm)
